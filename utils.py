import io
import re
import json
import pypdf
import docx
from groq import Groq
import prompts

def extract_text_from_file(uploaded_file):
    """
    Extracts text from uploaded file bytes (PDF, DOCX, TXT).
    """
    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()
    # Reset read pointer just in case
    uploaded_file.seek(0)
    
    if file_name.endswith('.pdf'):
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        text = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n".join(text)
        
    elif file_name.endswith('.docx'):
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                # Avoid empty cells cluttering
                text.append(" | ".join(row_text))
        return "\n".join(text)
        
    elif file_name.endswith('.txt'):
        return file_bytes.decode('utf-8', errors='ignore')
        
    else:
        raise ValueError("Unsupported file format. Please upload PDF, DOCX, or TXT.")

def clean_and_parse_json(text, fallback_type="analysis"):
    """
    Cleans Groq JSON output (removes markdown backticks if present) and parses it.
    """
    cleaned = text.strip()
    
    # Remove markdown code blocks if Llama formats it
    if cleaned.startswith("```json"):
        cleaned = cleaned[7:]
    elif cleaned.startswith("```"):
        cleaned = cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
        
    cleaned = cleaned.strip()
    
    # Try parsing
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        # Regex fallback to find the first '{' and last '}'
        match = re.search(r'\{.*\}', cleaned, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass
                
    # Ultimate fallback if it fails completely
    if fallback_type == "analysis":
        return {
            "document_type": "Legal Document",
            "parties": ["Could not be parsed"],
            "effective_date": "Not found",
            "expiration_date": "Not found",
            "governing_law": "Not found",
            "jurisdiction": "Not found",
            "summary": "We extracted the text, but the structured JSON analysis failed to parse. Below is the raw text analysis.",
            "key_clauses": [
                {
                    "clause_name": "Raw Model Output",
                    "clause_text_snippet": "JSON Parse Failure",
                    "summary": text
                }
            ]
        }
    else:
        return {
            "overall_risk_score": 50,
            "overall_assessment": "JSON Parse Failure. The model output could not be loaded into JSON. Here is the raw analysis: " + text[:500] + "...",
            "risk_categories": {
                "Liability & Indemnity": "Medium",
                "Termination & Auto-renewals": "Medium",
                "Intellectual Property": "Medium",
                "Payment & Financials": "Medium"
            },
            "risks": [
                {
                    "clause_type": "Formatting issue",
                    "risk_level": "Medium",
                    "description": "Model output could not be parsed to JSON.",
                    "implication": "Structured representation is unavailable.",
                    "mitigation_suggestion": "Review the full contract text manually or try running the analysis again."
                }
            ]
        }

def analyze_document(document_text, api_key, model="llama-3.3-70b-versatile"):
    """
    Call Groq to perform document analysis & metadata extraction.
    """
    client = Groq(api_key=api_key)
    # Truncate text if it is extremely long to fit context safely (e.g. max 50,000 characters)
    # Llama 3.3 has a 128k context, but let's keep it safe.
    truncated_text = document_text[:60000]
    
    formatted_prompt = prompts.ANALYSIS_PROMPT.replace("{document_text}", truncated_text)
    
    response = client.chat.completions.create(
        messages=[
            {"role": "system", "content": prompts.SYSTEM_PROMPT},
            {"role": "user", "content": formatted_prompt}
        ],
        model=model,
        temperature=0.1, # Low temperature for accurate metadata extraction
        response_format={"type": "json_object"} # Force JSON mode on Groq
    )
    
    output_text = response.choices[0].message.content
    return clean_and_parse_json(output_text, "analysis")

def assess_risks(document_text, api_key, model="llama-3.3-70b-versatile"):
    """
    Call Groq to evaluate risks and red flags.
    """
    client = Groq(api_key=api_key)
    truncated_text = document_text[:60000]
    
    formatted_prompt = prompts.RISK_PROMPT.replace("{document_text}", truncated_text)
    
    response = client.chat.completions.create(
        messages=[
            {"role": "system", "content": prompts.SYSTEM_PROMPT},
            {"role": "user", "content": formatted_prompt}
        ],
        model=model,
        temperature=0.1, # Low temperature for consistent risk assessment
        response_format={"type": "json_object"}
    )
    
    output_text = response.choices[0].message.content
    return clean_and_parse_json(output_text, "risk")

def answer_question(document_text, chat_history, question, api_key, model="llama-3.3-70b-versatile"):
    """
    Call Groq to answer a user question grounded in the document text.
    """
    client = Groq(api_key=api_key)
    truncated_text = document_text[:50000]
    
    # Format chat history
    history_str = ""
    for turn in chat_history[-6:]: # Keep last 3 exchanges (6 turns) for context
        role = "User" if turn["role"] == "user" else "Assistant"
        history_str += f"{role}: {turn['content']}\n"
        
    formatted_prompt = prompts.CHAT_PROMPT.format(
        document_text=truncated_text,
        chat_history=history_str,
        question=question
    )
    
    response = client.chat.completions.create(
        messages=[
            {"role": "system", "content": prompts.SYSTEM_PROMPT},
            {"role": "user", "content": formatted_prompt}
        ],
        model=model,
        temperature=0.3
    )
    
    return response.choices[0].message.content

def draft_contract(contract_type, effective_date, governing_law, parties_details, key_terms, api_key, model="llama-3.3-70b-versatile"):
    """
    Call Groq to draft a contract template.
    """
    client = Groq(api_key=api_key)
    
    formatted_prompt = prompts.DRAFTING_PROMPT.format(
        contract_type=contract_type,
        effective_date=effective_date,
        governing_law=governing_law,
        parties_details=parties_details,
        key_terms=key_terms
    )
    
    response = client.chat.completions.create(
        messages=[
            {"role": "system", "content": prompts.SYSTEM_PROMPT},
            {"role": "user", "content": formatted_prompt}
        ],
        model=model,
        temperature=0.5 # Slightly higher temp for drafting flexibility
    )
    
    return response.choices[0].message.content
