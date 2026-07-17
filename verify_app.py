import os
import io
import docx
from groq import Groq
import utils

def test_docx_extraction():
    print("Testing DOCX extraction...")
    doc = docx.Document()
    doc.add_paragraph("This is a test paragraph for the legal assistant.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Party A"
    table.cell(0, 1).text = "Apex Technologies"
    table.cell(1, 0).text = "Party B"
    table.cell(1, 1).text = "Summit Analytics"
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Mocking streamlit file class
    class MockUploadedFile:
        def __init__(self, name, bytes_data):
            self.name = name
            self.bytes_data = bytes_data
        def read(self):
            return self.bytes_data
        def seek(self, pos):
            pass
            
    uploaded_file = MockUploadedFile("test_contract.docx", file_stream.getvalue())
    text = utils.extract_text_from_file(uploaded_file)
    print("Extracted Text:\n", text)
    assert "test paragraph" in text.lower()
    assert "Apex Technologies" in text
    print("DOCX Extraction: PASS")

def test_txt_extraction():
    print("Testing TXT extraction...")
    content = b"This is a plain text legal agreement draft."
    class MockUploadedFile:
        def __init__(self, name, bytes_data):
            self.name = name
            self.bytes_data = bytes_data
        def read(self):
            return self.bytes_data
        def seek(self, pos):
            pass
            
    uploaded_file = MockUploadedFile("test_contract.txt", content)
    text = utils.extract_text_from_file(uploaded_file)
    print("Extracted Text:\n", text)
    assert "plain text" in text.lower()
    print("TXT Extraction: PASS")

def test_groq_import():
    print("Testing Groq import...")
    try:
        # Just ensure client can be instantiated (even with empty key, just check import)
        client = Groq(api_key="gsk_dummykey")
        print("Groq Client Initialization: PASS")
    except Exception as e:
        print("Groq Client Initialization: FAIL -", str(e))
        raise

if __name__ == "__main__":
    print("=== STARTING APP VERIFICATION ===")
    test_docx_extraction()
    print("-" * 30)
    test_txt_extraction()
    print("-" * 30)
    test_groq_import()
    print("=== ALL STATIC TESTS PASSED ===")
